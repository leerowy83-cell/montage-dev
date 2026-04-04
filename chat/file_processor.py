"""
MontageDev — File Processor
Handles: PDF, DOCX, XLSX, CSV, TXT, code files, ZIP archives, JSON, XML, YAML, images.
Returns extracted text or base64 for images.
"""
import base64
import io
import json
import os
import xml.dom.minidom
import zipfile
import csv

SUPPORTED_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".log", ".rst",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".htm",
    ".css", ".scss", ".sass", ".less",
    ".java", ".c", ".cpp", ".h", ".cs", ".go", ".rs",
    ".rb", ".php", ".swift", ".kt", ".scala", ".r",
    ".sh", ".bash", ".zsh", ".fish",
    ".sql", ".graphql", ".prisma",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".env",
    ".xml", ".svg",
    ".dockerfile", ".makefile",
    ".gitignore", ".editorconfig",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".avif", ".bmp"}
MAX_TEXT_CHARS = 80_000  # ~20k tokens


def process_file(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point. Returns:
    {
        "type":     "text" | "image",
        "filename": str,
        "content":  str,           # for text
        "base64":   str,           # for images
        "mimetype": str,           # for images
        "summary":  str,           # human-readable description
        "error":    str | None,
    }
    """
    ext = os.path.splitext(filename.lower())[1]
    result = {"filename": filename, "error": None}

    try:
        if ext == ".pdf":
            return _process_pdf(file_bytes, filename)
        elif ext in (".docx", ".doc"):
            return _process_docx(file_bytes, filename)
        elif ext in (".xlsx", ".xls"):
            return _process_xlsx(file_bytes, filename)
        elif ext == ".csv":
            return _process_csv(file_bytes, filename)
        elif ext in (".zip",):
            return _process_zip(file_bytes, filename)
        elif ext in IMAGE_EXTENSIONS:
            return _process_image(file_bytes, filename, ext)
        elif ext in SUPPORTED_TEXT_EXTENSIONS or ext == "":
            return _process_text(file_bytes, filename)
        else:
            # Try as text anyway
            return _process_text(file_bytes, filename)
    except Exception as e:
        return {
            "type": "error",
            "filename": filename,
            "content": "",
            "summary": f"Failed to process file: {e}",
            "error": str(e),
        }


# ──────────────────────────────────────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────────────────────────────────────

def _process_pdf(data: bytes, filename: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i+1} ---\n{text.strip()}")

    full_text = "\n\n".join(pages)
    if not full_text.strip():
        full_text = "[PDF contained no extractable text — may be a scanned image]"

    return {
        "type": "text",
        "filename": filename,
        "content": _truncate(full_text),
        "summary": f"PDF document, {len(reader.pages)} page(s)",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────────────────────

def _process_docx(data: bytes, filename: str) -> dict:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for i, table in enumerate(doc.tables):
        parts.append(f"\n[Table {i+1}]")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))

    full_text = "\n".join(parts)
    return {
        "type": "text",
        "filename": filename,
        "content": _truncate(full_text),
        "summary": f"Word document, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} table(s)",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# XLSX
# ──────────────────────────────────────────────────────────────────────────────

def _process_xlsx(data: bytes, filename: str) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"\n[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_str = " | ".join(str(c) if c is not None else "" for c in row)
            if row_str.strip(" |"):
                parts.append(row_str)
        if len(parts) > 5000:
            parts.append("... [truncated for length]")
            break

    return {
        "type": "text",
        "filename": filename,
        "content": _truncate("\n".join(parts)),
        "summary": f"Excel workbook with {len(wb.sheetnames)} sheet(s): {', '.join(wb.sheetnames)}",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# CSV
# ──────────────────────────────────────────────────────────────────────────────

def _process_csv(data: bytes, filename: str) -> dict:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    formatted = "\n".join(" | ".join(row) for row in rows[:500])
    if len(rows) > 500:
        formatted += f"\n... [{len(rows) - 500} more rows truncated]"

    return {
        "type": "text",
        "filename": filename,
        "content": _truncate(formatted),
        "summary": f"CSV file, {len(rows)} rows × {len(rows[0]) if rows else 0} columns",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# ZIP
# ──────────────────────────────────────────────────────────────────────────────

def _process_zip(data: bytes, filename: str) -> dict:
    parts = [f"ZIP archive: {filename}\n"]
    processed = 0

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
        parts.append(f"Contents ({len(names)} files):")
        for name in names:
            parts.append(f"  • {name}")
        parts.append("")

        # Extract and process readable files (skip large binaries)
        for name in names:
            if processed >= 20:
                parts.append(f"\n[Stopped after 20 files to avoid truncation]")
                break
            info = zf.getinfo(name)
            if info.is_dir() or info.file_size > 500_000:
                continue
            ext = os.path.splitext(name.lower())[1]
            if ext in IMAGE_EXTENSIONS:
                parts.append(f"\n--- {name} [image, not shown] ---")
                continue
            try:
                file_bytes = zf.read(name)
                inner = process_file(file_bytes, name)
                if inner["type"] == "text" and inner["content"]:
                    parts.append(f"\n--- {name} ---\n{inner['content'][:3000]}")
                    processed += 1
            except Exception:
                parts.append(f"\n--- {name} [could not read] ---")

    return {
        "type": "text",
        "filename": filename,
        "content": _truncate("\n".join(parts)),
        "summary": f"ZIP archive containing {len(names)} file(s)",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# IMAGE
# ──────────────────────────────────────────────────────────────────────────────

def _process_image(data: bytes, filename: str, ext: str) -> dict:
    mime_map = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".webp": "image/webp", ".avif": "image/avif",
        ".bmp": "image/bmp",
    }
    mimetype = mime_map.get(ext, "image/jpeg")
    b64 = base64.b64encode(data).decode("utf-8")

    return {
        "type": "image",
        "filename": filename,
        "content": "",
        "base64": b64,
        "mimetype": mimetype,
        "summary": f"Image file ({ext[1:].upper()})",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# PLAIN TEXT / CODE
# ──────────────────────────────────────────────────────────────────────────────

def _process_text(data: bytes, filename: str) -> dict:
    text = data.decode("utf-8", errors="replace")
    ext = os.path.splitext(filename.lower())[1]

    # Pretty-print JSON
    if ext == ".json":
        try:
            parsed = json.loads(text)
            text = json.dumps(parsed, indent=2)
        except Exception:
            pass

    # Pretty-print XML
    if ext == ".xml":
        try:
            dom = xml.dom.minidom.parseString(text)
            text = dom.toprettyxml(indent="  ")
        except Exception:
            pass

    lines = text.count("\n") + 1
    return {
        "type": "text",
        "filename": filename,
        "content": _truncate(text),
        "summary": f"Text/code file ({ext or 'unknown type'}), {lines} lines",
        "error": None,
    }


# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def _truncate(text: str) -> str:
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS] + "\n\n[... content truncated to fit context window ...]"
    return text
